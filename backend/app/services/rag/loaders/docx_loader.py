"""
DOCX document loader.

Extracts text from Microsoft Word documents.
"""

from typing import Optional, Dict, Any
import os

from app.services.rag.loaders.base import BaseDocumentLoader
from app.services.rag.protocols import Document


class DocxLoader(BaseDocumentLoader):
    """
    Loader for Microsoft Word (.docx) files.

    Uses python-docx to extract text from Word documents.
    """

    def supports(self, source: str) -> bool:
        """
        Check if this loader supports the given source.

        Args:
            source: File path

        Returns:
            True if file has .docx extension
        """
        return source.lower().endswith(".docx")

    async def load(
        self,
        source: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Document:
        """
        Load a DOCX document from file.

        Args:
            source: File path to DOCX
            metadata: Optional metadata to attach

        Returns:
            Loaded document with extracted text
        """
        from docx import Document as DocxDocument

        # Read DOCX
        doc = DocxDocument(source)

        # Extract text from all paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_texts.append(row_text)

        # Combine paragraphs and tables
        content = "\n\n".join(paragraphs)
        if table_texts:
            content += "\n\n" + "\n".join(table_texts)

        # Clean text
        content = self.clean_text(content)

        # Build metadata
        full_metadata = self.build_metadata(source, metadata)
        full_metadata["file_size"] = os.path.getsize(source)
        full_metadata["mime_type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # Extract document properties if available
        core_properties = doc.core_properties
        doc_info = {}

        if core_properties.title:
            doc_info["title"] = core_properties.title
        if core_properties.author:
            doc_info["author"] = core_properties.author
        if core_properties.subject:
            doc_info["subject"] = core_properties.subject
        if core_properties.created:
            doc_info["created"] = core_properties.created.isoformat()
        if core_properties.modified:
            doc_info["modified"] = core_properties.modified.isoformat()

        if doc_info:
            full_metadata["document_info"] = doc_info

        return Document(
            content=content,
            metadata=full_metadata,
        )
